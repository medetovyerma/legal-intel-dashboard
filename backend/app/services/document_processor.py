"""
Document processing service for extracting text from PDF and DOCX files
"""
import os
import logging
from typing import Tuple, Optional
from pathlib import Path

import PyPDF2
import pdfplumber
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing and extracting text from legal documents"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> Tuple[str, bool]:
        """
        Extract text from PDF file using multiple methods for robustness
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Tuple of (extracted_text, success_flag)
        """
        try:
            # First try with pdfplumber (better for complex layouts)
            text = DocumentProcessor._extract_with_pdfplumber(file_path)
            if text and len(text.strip()) > 50:  # Reasonable text threshold
                return text, True
            
            # Fallback to PyPDF2
            text = DocumentProcessor._extract_with_pypdf2(file_path)
            if text and len(text.strip()) > 50:
                return text, True
            
            logger.warning(f"Limited text extracted from PDF: {file_path}")
            return text or "", False
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            return "", False
    
    @staticmethod
    def _extract_with_pdfplumber(file_path: str) -> str:
        """Extract text using pdfplumber library"""
        text_content = []
        
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
        
        return "\n".join(text_content)
    
    @staticmethod
    def _extract_with_pypdf2(file_path: str) -> str:
        """Extract text using PyPDF2 library as fallback"""
        text_content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
        
        return "\n".join(text_content)
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> Tuple[str, bool]:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Tuple of (extracted_text, success_flag)
        """
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            extracted_text = "\n".join(text_content)
            success = len(extracted_text.strip()) > 10
            
            if not success:
                logger.warning(f"Limited text extracted from DOCX: {file_path}")
            
            return extracted_text, success
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            return "", False
    
    @staticmethod
    def process_document(file_path: str) -> Tuple[str, bool]:
        """
        Process document and extract text based on file extension
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple of (extracted_text, success_flag)
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return "", False
        
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return DocumentProcessor.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return DocumentProcessor.extract_text_from_docx(file_path)
        else:
            logger.error(f"Unsupported file format: {file_extension}")
            return "", False
    
    @staticmethod
    def validate_file(file_path: str, max_size: int) -> Tuple[bool, Optional[str]]:
        """
        Validate file size and format
        
        Args:
            file_path: Path to the file
            max_size: Maximum allowed file size in bytes
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                return False, "File not found"
            
            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > max_size:
                return False, f"File size ({file_size} bytes) exceeds maximum allowed size ({max_size} bytes)"
            
            # Check file extension
            file_extension = Path(file_path).suffix.lower()
            allowed_extensions = {'.pdf', '.docx'}
            if file_extension not in allowed_extensions:
                return False, f"Unsupported file format: {file_extension}. Allowed formats: {allowed_extensions}"
            
            return True, None
            
        except Exception as e:
            return False, f"Error validating file: {str(e)}"